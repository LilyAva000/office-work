from typing import Dict, List, Any
from docx import Document
from docx.shared import Mm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def set_font(cell, name: str = "仿宋_GB2312", size: int = 12) -> None:
    """设置单元格字体为仿宋_GB2312，字号size。"""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
            run.font.size = Pt(size)


def set_cell_bg_yellow(cell) -> None:
    """设置单元格背景为黄色。"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old_shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(old_shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'FFFF00')
    tcPr.append(shd)


def insert_photo(table, photo_path: str) -> None:
    """在表格中插入照片。"""
    for row in table.rows:
        for cell in row.cells:
            if '（照片' in cell.text or '(照片' in cell.text:
                cell.text = ''
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run()
                try:
                    run.add_picture(photo_path, width=Mm(35))
                except Exception as e:
                    print("插入图片失败：", e)
                return


def fill_basic_info(table, info: Dict[str, Any], family_header_row_idx: int) -> None:
    """填写基本信息区。"""
    for row_idx, row in enumerate(table.rows):
        if row_idx >= family_header_row_idx:
            break
        for idx, cell in enumerate(row.cells):
            key = cell.text.strip()
            if key in info and key not in ['个人简历', '家庭情况', '照片']:
                if idx + 1 < len(row.cells):
                    content_cell = row.cells[idx + 1]
                    content_cell.text = info[key]
                    set_font(content_cell)
                    set_cell_bg_yellow(content_cell)


def fill_contact_info(table, info: Dict[str, Any], family_header_row_idx: int) -> None:
    """填写通讯地址和手机。"""
    for row_idx, row in enumerate(table.rows):
        if row_idx >= family_header_row_idx:
            break
        if '通讯地址' in row.cells[0].text:
            row.cells[1].text = info.get('通讯地址', '')
            set_font(row.cells[1])
            set_cell_bg_yellow(row.cells[1])
            row.cells[-1].text = info.get('手机', '')
            set_font(row.cells[-1])
            set_cell_bg_yellow(row.cells[-1])


def fill_resume(table, info: Dict[str, Any], family_header_row_idx: int) -> None:
    """填写个人简历。"""
    for row_idx, row in enumerate(table.rows):
        if row_idx >= family_header_row_idx:
            break
        if '个人简历' in row.cells[0].text:
            row.cells[1].text = info.get('个人简历', '')
            set_font(row.cells[1])
            set_cell_bg_yellow(row.cells[1])


def fill_family_info(table, family_info: List[Dict[str, str]], family_header_row_idx: int) -> None:
    """根据表头动态填写家庭情况，自动适应表结构。"""
    # 获取表头行内容
    header_row = table.rows[family_header_row_idx]
    header_fields = [cell.text.strip() for cell in header_row.cells]
    for i, member in enumerate(family_info):
        data_row_idx = family_header_row_idx + 1 + i
        if data_row_idx < len(table.rows):
            data_row = table.rows[data_row_idx]
            for col_idx, header in enumerate(header_fields):
                if header and header in member and col_idx < len(data_row.cells):
                    data_row.cells[col_idx].text = member[header]
                    set_font(data_row.cells[col_idx])
                    set_cell_bg_yellow(data_row.cells[col_idx])


def find_family_header_row(table) -> int:
    """查找家庭情况表头行索引。"""
    for i, row in enumerate(table.rows):
        titles = [cell.text.strip() for cell in row.cells[:4]]
        if titles[:3] == ['关系', '姓名', '身份证号']:
            return i
    raise ValueError("未找到‘家庭情况’表头行（关系/姓名/身份证号）。")


def print_table_rows(table, decs="") -> None:
    """打印表格每一行的内容。"""
    print(f"--------------------{decs}--------------------")
    for idx, row in enumerate(table.rows):
        values = [cell.text for cell in row.cells]
        print(f"Row {idx}: {values}")


def main(template_path, output_path):
    doc = Document(template_path)
    table = doc.tables[0]
    print_table_rows(table, '填写前')


    info = {
        '姓名': '张三',
        '身份证号': '1234567890',
        '出生年月': '1990-01-01',
        '入党日期': '2016-02-15',
        '政治面貌': '团员',
        '籍贯': '北京',
        '文化程度': '本科',
        '性别': '男',
        '民族': '汉',
        '通讯地址': '北京市朝阳区',
        '手机': '13800138000',
        '个人简历': '2008-2012 就读于XX大学\n2012-2020 就职于XX公司',
        '家庭情况': [
            {'关系': '父亲', '姓名': '张爸', '身份证号': '1111111111', '工作单位': '某公司'},
            {'关系': '母亲', '姓名': '李妈', '身份证号': '2222222222', '工作单位': '某单位'},
        ],
        '照片': r'C:\Users\jfw\Pictures\img\loading.png',
        '有无出国（境）证件': '无',
        '出国（境）情况': '无',
    }
    family_header_row_idx = find_family_header_row(table)
    fill_basic_info(table, info, family_header_row_idx)
    fill_contact_info(table, info, family_header_row_idx)
    fill_resume(table, info, family_header_row_idx)
    if '照片' in info and info['照片']:
        insert_photo(table, info['照片'])
    fill_family_info(table, info.get('家庭情况', []), family_header_row_idx)
    print_table_rows(table, '填写后')
    doc.save(output_path)


if __name__ == '__main__':
    template_path = 'templates/word-table3.docx'
    output_path = 'output/robust-word-filled3.docx'
    main(template_path, output_path)