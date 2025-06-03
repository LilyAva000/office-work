from docx import Document
from docx.shared import Mm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def set_fangsong(cell, size=12):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = '仿宋_GB2312'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            run.font.size = Pt(size)

def set_cell_bg_yellow(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # 删除已有的shd，防止叠加
    for old_shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(old_shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'FFFF00')
    tcPr.append(shd)

doc = Document('templates/word-table.docx')
table = doc.tables[0]

info = {
    '姓名': '张三',
    '身份证号': '1234567890',
    '出生年月': '1990-01-01',
    '政治面貌': '团员',
    '籍贯': '北京',
    '文化程度': '本科',
    '性别': '男',
    '民族': '汉',
    '通讯地址': '北京市朝阳区',
    '手机': '13800138000',
    '个人简历': '2008-2012 就读于XX大学\n2012-2020 就职于XX公司',
    '家庭情况': [
        {'关系': '父亲', '姓名': '张爸', '身份证号': '1111111111', '工作单位': '某公司'},
        {'关系': '母亲', '姓名': '李妈', '身份证号': '2222222222', '工作单位': '某单位'},
    ],
    '照片': r'C:\Users\jfw\Pictures\img\loading.png'
}

# 找到“家庭情况”表头所在行
family_header_row_idx = None
for i, row in enumerate(table.rows):
    titles = [cell.text.strip() for cell in row.cells[:4]]
    # print(f"Row {i}: {titles}")  # 调试
    if titles[:3] == ['关系', '姓名', '身份证号']:
        family_header_row_idx = i
        break

if family_header_row_idx is None:
    raise ValueError("未能在表格中找到‘家庭情况’表头行（关系/姓名/身份证号）。请检查模板文件。")

# 基本信息填写（只到家庭情况表头前一行，严格不覆盖表头）
for row_idx, row in enumerate(table.rows):
    if row_idx >= family_header_row_idx:
        break
    for idx, cell in enumerate(row.cells):
        key = cell.text.strip()
        if key in info and key not in ['个人简历', '家庭情况', '照片']:
            if idx + 1 < len(row.cells):
                row.cells[idx + 1].text = info[key]
                set_fangsong(row.cells[idx + 1])
                set_cell_bg_yellow(row.cells[idx + 1])

# 通讯地址和手机
for row_idx, row in enumerate(table.rows):
    if row_idx >= family_header_row_idx:
        break
    if '通讯地址' in row.cells[0].text:
        row.cells[1].text = info['通讯地址']
        set_fangsong(row.cells[1])
        set_cell_bg_yellow(row.cells[1])
        row.cells[-1].text = info['手机']
        set_fangsong(row.cells[-1])
        set_cell_bg_yellow(row.cells[-1])

# 个人简历
for row_idx, row in enumerate(table.rows):
    if row_idx >= family_header_row_idx:
        break
    if '个人简历' in row.cells[0].text:
        row.cells[1].text = info['个人简历']
        set_fangsong(row.cells[1])
        set_cell_bg_yellow(row.cells[1])

# 插入照片（不建议背景色，否则图片会遮挡色块无法看清）
if '照片' in info and info['照片']:
    for row in table.rows:
        for cell in row.cells:
            if '（照片）' in cell.text:
                cell.text = ''
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = paragraph.add_run()
                try:
                    run.add_picture(info['照片'], width=Mm(35))
                except Exception as e:
                    print("插入图片失败：", e)
                break
        else:
            continue
        break

# 家庭情况（只写数据行，不碰表头）
family_info = info['家庭情况']
# 数据行从family_header_row_idx+1开始
for i, member in enumerate(family_info):
    data_row_idx = family_header_row_idx + 1 + i
    if data_row_idx < len(table.rows):
        data_row = table.rows[data_row_idx]
        cells = data_row.cells
        if len(cells) >= 7:
            cells[0].text = member.get('关系', '')
            set_fangsong(cells[0])
            set_cell_bg_yellow(cells[0])
            cells[1].text = member.get('姓名', '')
            set_fangsong(cells[1])
            set_cell_bg_yellow(cells[1])
            cells[2].text = member.get('身份证号', '')
            set_fangsong(cells[2])
            set_cell_bg_yellow(cells[2])
            cells[5].text = member.get('工作单位', '')
            set_fangsong(cells[5])
            set_cell_bg_yellow(cells[5])

doc.save('output/word-table5.docx')