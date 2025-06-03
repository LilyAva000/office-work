from docx import Document
from docx.shared import Cm
import json
import os
import requests

OLLAMA_API_URL = "http://12.191.203.213:11434/api/generate"
OLLAMA_MODEL = "qwen2.5-coder:14b"  # 你可以根据实际ollama模型名称修改

def is_pic_path(value):
    return isinstance(value, str) and value.lower().endswith(('.png', '.jpg', '.jpeg'))

def insert_photo_to_cell(cell, img_path):
    if os.path.exists(img_path):
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(img_path, width=Cm(2.5))
    else:
        cell.text = "图片未找到"

def get_table_merge_main_cells(doc):
    merge_main_cells = set()
    for t_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                tc = cell._tc
                gridspan = tc.xpath('.//w:gridSpan')
                vmerge = tc.xpath('.//w:vMerge')
                if gridspan or any(x.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '') == 'restart' for x in vmerge):
                    merge_main_cells.add((t_idx, row_idx, col_idx))
                elif not vmerge:
                    merge_main_cells.add((t_idx, row_idx, col_idx))
    return merge_main_cells

def fill_word_with_llm_json(filepath, llm_result_json, output_path):
    doc = Document(filepath)
    llm_result = json.loads(llm_result_json)
    merge_main_cells = get_table_merge_main_cells(doc)
    for item in llm_result:
        tidx, row, col, value = item["table_index"], item["row"], item["col"], item["value"]
        if (tidx, row, col) in merge_main_cells:
            cell = doc.tables[tidx].rows[row].cells[col]
            if is_pic_path(value):
                insert_photo_to_cell(cell, value)
            else:
                cell.text = str(value)
    doc.save(output_path)

def extract_word_structure(filepath):
    doc = Document(filepath)
    structure = []
    merge_main_cells = get_table_merge_main_cells(doc)
    for ti, table in enumerate(doc.tables):
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if (ti, i, j) not in merge_main_cells:
                    continue
                info = {
                    "table_index": ti,
                    "row": i,
                    "col": j,
                    "value": cell.text.strip(),
                    "left": row.cells[j-1].text.strip() if j > 0 else None,
                    "up": table.rows[i-1].cells[j].text.strip() if i > 0 else None,
                }
                if not info["value"] or "{" in info["value"]:
                    structure.append(info)
    return structure

def build_prompt(structure):
    info = {
        '姓名': '张三',
        '身份证号': '1234567890',
        '出生年月': '1990-01-01',
        '入党日期': '2016-02-15',
        '政治面貌': '团员',
        '籍贯': '北京',
        '文化程度': '本科',
        '性别': '男',
        '民族': '汉',
        '通讯地址': '北京市朝阳区',
        '手机': '13800138000',
        '个人简历': '2008-2012 就读于XX大学\n2012-2020 就职于XX公司',
        '家庭情况': [
            {'关系': '父亲', '姓名': '张爸', '身份证号': '1111111111', '工作单位': '某公司'},
            {'关系': '母亲', '姓名': '李妈', '身份证号': '2222222222', '工作单位': '某单位'},
        ],
        '照片': r'C:\Users\jfw\Pictures\img\loading.png',
        '有无出国（境）证件': '无',
        '出国（境）情况': '无',
    }
    prompt = (
        f"请根据以下表格结构信息，针对每一个待填项，根据上下文与字段名，给出应该填写的内容。\n"
        f"内容是{info}\n\n"
        f"输出格式为JSON，结构信息如下：\n\n\n"
        f"{json.dumps(structure, ensure_ascii=False, indent=2)}\n"
        f"请填写所有空白项。"
    )
    return prompt

def call_ollama(prompt, model=OLLAMA_MODEL):
    headers = {"Content-Type": "application/json"}
    data = {
        "model": model,
        "prompt": prompt,
        "stream": False
    }
    resp = requests.post(OLLAMA_API_URL, headers=headers, data=json.dumps(data), timeout=300)
    resp.raise_for_status()
    # ollama返回内容格式: {'response': '....'}
    # 有的模型返回的JSON内容前后有多余文本，需要清理
    content = resp.json().get("response", "")
    # 尝试提取有效的JSON片段
    try:
        start = content.index('[')
        end = content.rindex(']') + 1
        json_str = content[start:end]
        return json_str
    except Exception:
        # 若无法提取，直接返回
        return content

def fill_word_auto():
    structure = extract_word_structure("templates/word-table3.docx")
    prompt = build_prompt(structure)
    print("正在调用ollama大模型，请稍候...")
    llm_result_json = call_ollama(prompt)
    with open('llm.json', 'w', encoding='utf-8') as f:
        f.write(llm_result_json)
    fill_word_with_llm_json("templates/word-table3.docx", llm_result_json, "word-table-filled.docx")
    print("自动填表完成，已保存为word-table-filled.docx")

if __name__ == "__main__":
    fill_word_auto()