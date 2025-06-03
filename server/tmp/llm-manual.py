from docx import Document
from docx.shared import Cm
import json
import os

def is_pic_path(value):
    return isinstance(value, str) and value.lower().endswith(('.png', '.jpg', '.jpeg'))

def insert_photo_to_cell(cell, img_path):
    if os.path.exists(img_path):
        # 清空文字
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(img_path, width=Cm(2.5))  # 可根据表格实际宽度调整
    else:
        cell.text = "图片未找到"

def get_table_merge_main_cells(doc):
    """获取每个表格的合并主格坐标集合，{(table_index, row, col): True}"""
    merge_main_cells = set()
    for t_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                tc = cell._tc
                gridspan = tc.xpath('.//w:gridSpan')
                vmerge = tc.xpath('.//w:vMerge')
                # 合并主格：无vMerge或vMerge为restart
                if gridspan or any(x.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '') == 'restart' for x in vmerge):
                    merge_main_cells.add((t_idx, row_idx, col_idx))
                # 不是合并主格但为普通单元格也要加入
                elif not vmerge:
                    merge_main_cells.add((t_idx, row_idx, col_idx))
    return merge_main_cells

def fill_word_with_llm_json(filepath, llm_result_json, output_path):
    doc = Document(filepath)
    llm_result = json.loads(llm_result_json)
    merge_main_cells = get_table_merge_main_cells(doc)
    for item in llm_result:
        tidx, row, col, value = item["table_index"], item["row"], item["col"], item["value"]
        # 只填合并主格或普通格
        if (tidx, row, col) in merge_main_cells:
            cell = doc.tables[tidx].rows[row].cells[col]
            if is_pic_path(value):
                insert_photo_to_cell(cell, value)
            else:
                cell.text = str(value)
        # 跳过非主格，避免覆盖
    doc.save(output_path)

def extract_word_structure(filepath):
    doc = Document(filepath)
    structure = []
    merge_main_cells = get_table_merge_main_cells(doc)
    for ti, table in enumerate(doc.tables):
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                # 只输出主格
                if (ti, i, j) not in merge_main_cells:
                    continue
                info = {
                    "table_index": ti,
                    "row": i,
                    "col": j,
                    "value": cell.text.strip(),
                    "left": row.cells[j-1].text.strip() if j > 0 else None,
                    "up": table.rows[i-1].cells[j].text.strip() if i > 0 else None,
                }
                # 空白或模板变量
                if not info["value"] or "{" in info["value"]:
                    structure.append(info)
    return structure

def build_prompt(structure):
    info = {
        '姓名': '张三',
        '身份证号': '1234567890',
        '出生年月': '1990-01-01',
        '入党日期': '2016-02-15',
        '政治面貌': '团员',
        '籍贯': '北京',
        '文化程度': '本科',
        '性别': '男',
        '民族': '汉',
        '通讯地址': '北京市朝阳区',
        '手机': '13800138000',
        '个人简历': '2008-2012 就读于XX大学\n2012-2020 就职于XX公司',
        '家庭情况': [
            {'关系': '父亲', '姓名': '张爸', '身份证号': '1111111111', '工作单位': '某公司'},
            {'关系': '母亲', '姓名': '李妈', '身份证号': '2222222222', '工作单位': '某单位'},
        ],
        '照片': r'C:\Users\jfw\Pictures\img\loading.png',
        '有无出国（境）证件': '无',
        '出国（境）情况': '无',
    }
    prompt = (
        f"请根据以下表格结构信息，针对每一个待填项，根据上下文与字段名，给出应该填写的内容。\n"
        f"内容是{info}\n\n"
        f"输出格式为JSON，结构信息如下：\n\n\n"
        f"{json.dumps(structure, ensure_ascii=False, indent=2)}\n"
        f"请填写所有空白项。"
    )
    return prompt

def get_llm_prompt():
    struct = extract_word_structure("templates/word-table3.docx")
    prompt = build_prompt(struct)
    print(prompt)
    return prompt

def fill_word():
    with open('llm.json', 'r', encoding='utf-8') as f:
        llm_result_json = json.load(f)
    fill_word_with_llm_json("templates/word-table3.docx", json.dumps(llm_result_json, ensure_ascii=False), "word-table-filled.docx")

if __name__ == "__main__":
    # get_llm_prompt()
    fill_word()